import streamlit as st
import openai
import os
from dotenv import load_dotenv
from docx import Document
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit
from full_resume_template import build_resume_docx
import json
from datetime import date
from docx.shared import Pt
import math

# Load environment variables from .env
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

st.set_page_config(page_title="Personalized Resume & Cover Letter Generator", page_icon="🧑‍💼", layout="centered")
st.title("🧑‍💼 Personalized Resume & Cover Letter Generator")

# Initialize session state for stepper and data
if 'step' not in st.session_state:
    st.session_state.step = 1
if 'job_title' not in st.session_state:
    st.session_state.job_title = ''
if 'job_description' not in st.session_state:
    st.session_state.job_description = ''
if 'resume_content' not in st.session_state:
    st.session_state.resume_content = ''
if 'polished_resume' not in st.session_state:
    st.session_state.polished_resume = ''
if 'polish_prompt' not in st.session_state:
    st.session_state.polish_prompt = ''
if 'company_name' not in st.session_state:
    st.session_state.company_name = ''
if 'cover_letter' not in st.session_state:
    st.session_state.cover_letter = ''
if 'preferences' not in st.session_state:
    st.session_state.preferences = ''
if 'placement_scores' not in st.session_state:
    st.session_state.placement_scores = ''
if 'uploaded_resume_bytes' not in st.session_state:
    st.session_state.uploaded_resume_bytes = None

# Stepper navigation
st.sidebar.title("Navigation")
step = st.sidebar.radio(
    "Go to step:",
    ["1. Resume Polisher", "2. Cover Letter Generator", "3. Placement Scores"],
    index=st.session_state.step-1
)
st.session_state.step = int(step[0])

# Step 1: Resume Polisher
if st.session_state.step == 1:
    st.header("📝 Step 1: Resume Polisher (GPT-4)")
    st.write("Upload your resume as a DOCX file, or paste/edit the content below. The extracted text will be used for polishing.")
    uploaded_file = st.file_uploader("Upload Resume (DOCX)", type=["docx"])
    resume_text = ""
    if uploaded_file is not None:
        st.session_state.uploaded_resume_bytes = uploaded_file.read()
        doc = Document(BytesIO(st.session_state.uploaded_resume_bytes))
        resume_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        st.session_state.resume_content = resume_text
        st.success("Resume loaded from DOCX. You can review and edit below.")
    with st.form("resume_polish_form"):
        job_title = st.text_input("Job Title", value=st.session_state.job_title, placeholder="Enter the job title...", key="job_title_input")
        company_name = st.text_input("Company Name", value=st.session_state.company_name, placeholder="Enter the company name you are applying to...", key="company_name_input")
        job_description = st.text_area("Job Description", value=st.session_state.job_description, placeholder="Paste the job description here...", height=120, key="job_description_input")
        resume_content = st.text_area("Resume Content", value=st.session_state.resume_content, placeholder="Paste your resume content here...", height=300, key="resume_content_input")
        polish_prompt = st.text_area("Polish Instruction (Optional)", value=st.session_state.polish_prompt, placeholder="Enter specific instructions or areas for improvement (optional)...", height=70, key="polish_prompt_input")
        submit = st.form_submit_button("Polish Resume")
    if submit:
        if not job_title or not job_description or not resume_content:
            st.error("Please provide the job title, job description, and your resume content.")
        else:
            with st.spinner("Polishing your resume with GPT-4..."):
                prompt = f"""
You are a helpful assistant that only returns valid JSON. 
Given the following resume content and job description, polish and update the resume to better fit the job. 
Return ONLY a valid JSON object matching this schema (no commentary, no markdown, no extra text):

{{
  "header": {{
    "name": "...",
    "email": "...",
    "phone": "...",
    "linkedin": "...",
    "github": "...",
    "website": "..."
  }},
  "summary": "...",
  "skills": [
    {{"category": "...", "details": "..."}}
  ],
  "experience": [
    {{
      "position": "...",
      "date_range": "...",
      "company": "...",
      "location": "...",
      "extra_line": "...",
      "applications": [
        {{"title": "...", "details": ["...", "..."]}}
      ]
    }}
  ],
  "education": {{
    "certificates": [{{"name": "...", "date": "..."}}],
    "specializations": [{{"institution": "...", "location": "...", "specialization": "...", "date": "..."}}],
    "degrees": [{{"university": "...", "location": "...", "date": "...", "degree": "..."}}]
  }}
}}

Resume Content:
{resume_content}

Job Title: {job_title}
Job Description: {job_description}
"""
                try:
                    response = openai.chat.completions.create(
                        model="gpt-4",
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.7,
                        max_tokens=1500,
                    )
                    gpt_output = response.choices[0].message.content.strip()
                    # Try to parse JSON
                    try:
                        resume_data = json.loads(gpt_output)
                    except Exception as json_err:
                        st.error(f"Could not parse resume JSON from GPT output.\nGPT output was:\n{gpt_output}")
                        raise Exception("JSON parsing failed")
                    st.session_state.job_title = job_title
                    st.session_state.job_description = job_description
                    st.session_state.resume_content = resume_content
                    st.session_state.polish_prompt = polish_prompt
                    st.session_state.polished_resume = gpt_output
                    st.session_state.resume_json_data = resume_data  # Save for editing
                    st.session_state.company_name = company_name
                    st.success("Resume polished! You can now review and edit each section below.")
                except Exception as e:
                    st.error(f"Error: {e}")
    if st.session_state.polished_resume:
        if st.button("Next: Cover Letter Generator"):
            st.session_state.step = 2
            st.rerun()

# Step 2: Cover Letter Generator
elif st.session_state.step == 2:
    st.header("📄 Step 2: Cover Letter Generator (GPT-4)")
    with st.form("cover_letter_form"):
        company_name = st.text_input("Company Name", value=st.session_state.company_name, placeholder="Enter the name of the company...", key="company_name_input")
        position_name = st.text_input("Position Name", value=st.session_state.job_title, placeholder="Enter the name of the position...", key="position_name_input")
        recruiter_name = st.text_input("Recruiter/Job Poster Name (Optional)", placeholder="Enter the recruiter's or job poster's name...", key="recruiter_name_input")
        job_description = st.text_area("Job Description Information", value=st.session_state.job_description, placeholder="Paste the job description here...", height=150, key="job_description_input")
        resume_content = st.text_area("Resume Content", value=st.session_state.polished_resume, placeholder="Paste your resume content here...", height=150, key="resume_content_input")
        submit = st.form_submit_button("Generate Cover Letter")
    if submit:
        if not company_name or not position_name or not job_description or not resume_content:
            st.error("Please provide all required fields.")
        else:
            with st.spinner("Generating your cover letter with GPT-4..."):
                if recruiter_name.strip():
                    salutation = f"Dear {recruiter_name},"
                else:
                    salutation = "Dear Hiring Manager,"
                prompt = (
                    f"Write a cover letter addressed as follows: {salutation}\n"
                    f"Generate a customized cover letter using the company name: {company_name}, the position applied for: {position_name}, and the job description: {job_description}. "
                    f"Ensure the cover letter highlights my qualifications and experience as detailed in the resume content: {resume_content}. "
                    "Adapt the content carefully to avoid including experiences not present in my resume but mentioned in the job description. "
                    "The goal is to emphasize the alignment between my existing skills and the requirements of the role."
                )
                try:
                    response = openai.chat.completions.create(
                        model="gpt-4",
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.7,
                        max_tokens=800,
                    )
                    cover_letter = response.choices[0].message.content.strip()
                    st.session_state.company_name = company_name
                    st.session_state.cover_letter = cover_letter
                    st.success("Cover letter generated! Proceed to the next step.")
                    st.text_area("Generated Cover Letter", value=cover_letter, height=300)
                    # Download options for cover letter
                    docx_buffer = BytesIO()
                    doc_out = Document()
                    # --- New: One-page limit logic ---
                    max_lines = 60
                    used_lines = 0.0
                    for para in [p for p in cover_letter.split('\n') if p.strip()]:
                        if used_lines + 1.5 > max_lines:
                            break
                        p = doc_out.add_paragraph(para)
                        p.paragraph_format.space_after = Pt(5)
                        used_lines += 1.5  # 1 line for text, 0.5 for 5pt after-space
                    doc_out.save(docx_buffer)
                    docx_buffer.seek(0)
                    st.session_state.cover_letter_docx_buffer = docx_buffer.getvalue()
                    st.download_button(
                        label="Download Cover Letter (DOCX)",
                        data=st.session_state.cover_letter_docx_buffer,
                        file_name="Cover_Letter.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="cover-letter-docx-gen"
                    )
                    # PDF download using reportlab (with word wrap)
                    pdf_buffer = BytesIO()
                    c = canvas.Canvas(pdf_buffer, pagesize=letter)
                    width, height = letter
                    left_margin = 40
                    right_margin = 40
                    max_width = width - left_margin - right_margin
                    y = height - 40
                    for line in cover_letter.split('\n'):
                        wrapped_lines = simpleSplit(line, 'Helvetica', 12, max_width)
                        for wline in wrapped_lines:
                            c.drawString(left_margin, y, wline)
                            y -= 18
                            if y < 40:
                                c.showPage()
                                y = height - 40
                    c.save()
                    pdf_buffer.seek(0)
                    st.session_state.cover_letter_pdf_buffer = pdf_buffer.getvalue()
                    st.download_button(
                        label="Download Cover Letter (PDF)",
                        data=st.session_state.cover_letter_pdf_buffer,
                        file_name="Cover_Letter.pdf",
                        mime="application/pdf",
                        key="cover-letter-pdf-gen"
                    )
                except Exception as e:
                    st.error(f"Error: {e}")
    if st.session_state.cover_letter:
        if st.button("Next: Placement Scores"):
            st.session_state.step = 3
            st.rerun()

# Step 3: Placement Scores
elif st.session_state.step == 3:
    st.header("📊 Step 3: Placement Scores (GPT-4)")
    with st.form("placement_scores_form"):
        job_title = st.text_input("Job Title", value=st.session_state.job_title, placeholder="Enter the job title...", key="job_title_input")
        job_description = st.text_area("Job Description", value=st.session_state.job_description, placeholder="Paste the job description here...", height=120, key="job_description_input")
        resume_content = st.text_area("Resume Content", value=st.session_state.polished_resume, placeholder="Paste your resume content here...", height=120, key="resume_content_input")
        preferences = st.text_area("Your Preferences/Goals (Optional)", value=st.session_state.preferences, placeholder="E.g. work-life balance, remote work, growth opportunities...", height=70, key="preferences_input")
        submit = st.form_submit_button("Get Placement Scores")
    if submit:
        if not job_title or not job_description or not resume_content:
            st.error("Please provide the job title, job description, and your resume content.")
        else:
            with st.spinner("Calculating placement scores with GPT-4..."):
                prompt = (
                    f"You are an honest career advisor. Given the following job title: {job_title}, job description: {job_description}, and resume: {resume_content}, "
                    f"analyze and provide two scores (0-100):\n"
                    f"1. How well does the candidate fit this job?\n"
                    f"2. How well does this job fit the candidate?\n"
                    f"For the second score, also consider the candidate's preferences/goals: {preferences if preferences.strip() else 'N/A'}.\n"
                    "For each score, provide a brief explanation. Be honest and do not hallucinate qualifications or experiences. Format your response as follows:\n"
                    "Fit for Job: <score>/100 - <explanation>\nJob Fit for You: <score>/100 - <explanation>"
                )
                try:
                    response = openai.chat.completions.create(
                        model="gpt-4",
                        messages=[{"role": "user", "content": prompt}],
                        temperature=0.2,
                        max_tokens=600,
                    )
                    result = response.choices[0].message.content.strip()
                    st.session_state.preferences = preferences
                    st.session_state.placement_scores = result
                    st.success("Placement scores generated!")
                    st.text_area("Placement Scores", value=result, height=200)
                except Exception as e:
                    st.error(f"Error: {e}")
    if st.session_state.placement_scores:
        if st.button("Back to Start"):
            st.session_state.step = 1
            st.session_state.polished_resume = ''
            st.session_state.cover_letter = ''
            st.session_state.placement_scores = ''
            st.rerun()

# Make sure download buttons are always available if buffers exist
if 'resume_docx_buffer' in st.session_state and st.session_state.resume_docx_buffer:
    name = st.session_state.resume_json_data.get('header', {}).get('name', 'Resume')
    company = st.session_state.company_name
    today_str = date.today().isoformat()
    resume_filename = f"{name} Resume for {company} ({today_str}).docx"
    st.download_button(
        label="Download Polished Resume (DOCX)",
        data=st.session_state.resume_docx_buffer,
        file_name=resume_filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="resume-docx-polish"
    )
if 'cover_letter_docx_buffer' in st.session_state and st.session_state.cover_letter_docx_buffer:
    cover_letter_filename = f"{name} Cover Letter for {company} ({today_str}).docx"
    st.download_button(
        label="Download Cover Letter (DOCX)",
        data=st.session_state.cover_letter_docx_buffer,
        file_name=cover_letter_filename,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="cover-letter-docx-gen"
    )
if 'cover_letter_pdf_buffer' in st.session_state and st.session_state.cover_letter_pdf_buffer:
    st.download_button(
        label="Download Cover Letter (PDF)",
        data=st.session_state.cover_letter_pdf_buffer,
        file_name="Cover_Letter.pdf",
        mime="application/pdf",
        key="cover-letter-pdf-gen"
    )

# --- Section Editing UI (all on one page) ---
if 'resume_json_data' in st.session_state and st.session_state.resume_json_data:
    st.header("✏️ Edit Your Resume Sections")
    resume_data = st.session_state.resume_json_data
    # Header
    st.subheader("Header")
    header = resume_data.get("header", {})
    header["name"] = st.text_input("Name", value=header.get("name", ""), key="header_name_edit_main")
    header["email"] = st.text_input("Email", value=header.get("email", ""), key="header_email_edit_main")
    header["phone"] = st.text_input("Phone", value=header.get("phone", ""), key="header_phone_edit_main")
    header["linkedin"] = st.text_input("LinkedIn", value=header.get("linkedin", ""), key="header_linkedin_edit_main")
    header["github"] = st.text_input("Github", value=header.get("github", ""), key="header_github_edit_main")
    header["website"] = st.text_input("Website", value=header.get("website", ""), key="header_website_edit_main")
    # Summary
    st.subheader("Professional Summary")
    resume_data["summary"] = st.text_area("Summary", value=resume_data.get("summary", ""), height=80)
    # Add creativity slider and regenerate button
    summary_temp = st.slider("Creativity (Temperature)", min_value=0.0, max_value=1.0, value=0.7, step=0.05, key="summary_temp_input")
    if st.button("Regenerate Professional Summary"):
        # Use current resume data as context
        skills = resume_data.get("skills", [])
        experience = resume_data.get("experience", [])
        education = resume_data.get("education", {})
        # Build a prompt
        prompt = (
            "Given the following resume data, generate a concise, impactful professional summary paragraph. "
            "Highlight the candidate's strengths, relevant experience, and technical skills. Return only the summary paragraph.\n"
            f"Skills: {skills}\nExperience: {experience}\nEducation: {education}"
        )
        try:
            response = openai.chat.completions.create(
                model="gpt-4",
                messages=[{"role": "user", "content": prompt}],
                temperature=summary_temp,
                max_tokens=200,
            )
            new_summary = response.choices[0].message.content.strip()
            resume_data["summary"] = new_summary
            st.success("Professional Summary regenerated!")
            st.rerun()
        except Exception as e:
            st.error(f"Error regenerating summary: {e}")
    # Skills
    st.subheader("Technical Skills")
    skills = resume_data.get("skills", [])
    skills = st.data_editor(skills, key="skills_editor_input", num_rows="dynamic")
    resume_data["skills"] = skills
    # Experience
    st.subheader("Professional Experience")
    experience = resume_data.get("experience", [])
    new_experience = []
    for i, job in enumerate(experience):
        # --- Pre-populate bullet_points from first application if needed (for jobs 2/3) ---
        if (not job.get("bullet_points")) and job.get("applications") and job["applications"] and job["applications"][0].get("title", "") == "":
            job["bullet_points"] = job["applications"][0].get("details", [])
            job["applications"] = job["applications"][1:]
        st.markdown(f"**Job {i+1}**")
        job["position"] = st.text_input(f"Position {i+1}", value=job.get("position", ""), key=f"pos_{i}_input")
        job["company"] = st.text_input(f"Company {i+1}", value=job.get("company", ""), key=f"comp_{i}_input")
        job["location"] = st.text_input(f"Location {i+1}", value=job.get("location", ""), key=f"loc_{i}_input")
        job["date_range"] = st.text_input(f"Date Range {i+1}", value=job.get("date_range", ""), key=f"date_{i}_input")
        # Dynamic bullet points (for jobs 2/3)
        bullet_points = job.get("bullet_points", [])
        if f"add_bullet_{i}_input" not in st.session_state:
            st.session_state[f"add_bullet_{i}_input"] = False
        st.markdown("**Bullet Points**")
        # Generate bullet points button
        if st.button("Generate Bullet Points", key=f"gen_bp_{i}_input"):
            job_prompt = f"Generate 3-5 concise, impactful bullet points describing the key responsibilities and achievements for the following job. Return ONLY a JSON list of strings, no commentary.\nJob Title: {job.get('position','')}\nCompany: {job.get('company','')}\nLocation: {job.get('location','')}\nDate Range: {job.get('date_range','')}"
            try:
                response = openai.chat.completions.create(
                    model="gpt-4",
                    messages=[{"role": "user", "content": job_prompt}],
                    temperature=0.7,
                    max_tokens=300,
                )
                bp_output = response.choices[0].message.content.strip()
                import json as _json
                try:
                    new_bullets = _json.loads(bp_output)
                    if isinstance(new_bullets, list):
                        bullet_points = new_bullets
                        job["bullet_points"] = bullet_points
                        st.session_state[f"bp_regen_{i}_input"] = bullet_points
                        st.success("Bullet points generated!")
                        st.rerun()
                    else:
                        st.error("OpenAI did not return a JSON list. Output was:\n" + bp_output)
                except Exception:
                    st.error("Could not parse bullet points JSON. Output was:\n" + bp_output)
            except Exception as e:
                st.error(f"Error generating bullet points: {e}")
        if f"bp_regen_{i}_input" in st.session_state:
            bullet_points = st.session_state[f"bp_regen_{i}_input"]
            job["bullet_points"] = bullet_points
        for b_idx, bullet in enumerate(bullet_points):
            cols = st.columns([0.85, 0.15])
            bullet_points[b_idx] = cols[0].text_input(f"Bullet Point {b_idx+1} (Job {i+1})", value=bullet, key=f"bp_{i}_input_{b_idx}")
            if cols[1].button("Remove", key=f"remove_bp_{i}_input_{b_idx}"):
                bullet_points.pop(b_idx)
                st.rerun()
        if st.button("Add Bullet Point", key=f"add_bp_{i}_input"):
            bullet_points.append("")
            st.rerun()
        job["bullet_points"] = bullet_points
        # Applications (for job 1 and any other jobs with titled applications)
        applications = job.get("applications", [])
        new_applications = []
        for j, app in enumerate(applications):
            app["title"] = st.text_input(f"Title {i+1}.{j+1}", value=app.get("title", ""), key=f"title_{i}_input_{j}")
            details_str = "\n".join(app.get("details", []))
            details_str = st.text_area(f"Details {i+1}.{j+1} (one bullet per line)", value=details_str, key=f"details_{i}_input_{j}")
            app["details"] = [line.strip() for line in details_str.splitlines() if line.strip()]
            new_applications.append(app)
        job["applications"] = new_applications
        new_experience.append(job)
    resume_data["experience"] = new_experience
    # Education
    st.subheader("Education")
    education = resume_data.get("education", {})
    certificates = education.get("certificates", [])
    certificates = st.data_editor(certificates, key="certificates_editor_input", num_rows="dynamic")
    education["certificates"] = certificates
    specializations = education.get("specializations", [])
    specializations = st.data_editor(specializations, key="specializations_editor_input", num_rows="dynamic")
    education["specializations"] = specializations
    degrees = education.get("degrees", [])
    degrees = st.data_editor(degrees, key="degrees_editor_input", num_rows="dynamic")
    education["degrees"] = degrees
    resume_data["education"] = education
    # Button to generate DOCX from edited data
    if st.button("Generate Resume with Edits"):
        docx_buffer = BytesIO()
        doc_out = build_resume_docx(resume_data)
        doc_out.save(docx_buffer)
        docx_buffer.seek(0)
        st.session_state.resume_docx_buffer = docx_buffer.getvalue()
        # Ensure filename variables are defined here
        name = resume_data.get('header', {}).get('name', 'Resume')
        company = st.session_state.company_name
        today_str = date.today().isoformat()
        resume_filename = f"{name} Resume for {company} ({today_str}).docx"
        st.download_button(
            label="Download Edited Resume (DOCX)",
            data=st.session_state.resume_docx_buffer,
            file_name=resume_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="resume-docx-edited"
        )

# --- New: Job Post Analysis and Line Allocation ---
def analyze_job_post_and_allocate_lines(job_description, resume_data, total_lines=60):
    """
    Use GPT-4 to analyze the job post and allocate lines to each section based on relevance.
    Returns a dict: {section: allocated_lines}
    """
    # Default priorities (can be tuned)
    default_priorities = {
        'summary': 1,
        'skills': 2,
        'experience': 5,
        'education': 3,
        'certificates': 1,
        'specializations': 1,
        'degrees': 1,
    }
    # Compose a prompt for GPT-4 to suggest priorities
    prompt = f"""
Given the following job description and resume data, assign a priority score (1-5, 5=most important) to each section: summary, skills, experience, education. 
Within education, also score: certificates, specializations, degrees. 
Return ONLY a valid JSON object like: {{'summary': 2, 'skills': 3, 'experience': 5, 'education': 3, 'certificates': 2, 'specializations': 2, 'degrees': 3}}

Job Description:\n{job_description}\nResume Data:\n{resume_data}
"""
    try:
        response = openai.chat.completions.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.2,
            max_tokens=200,
        )
        priorities = json.loads(response.choices[0].message.content.strip())
    except Exception:
        priorities = default_priorities
    # Normalize priorities to allocate total_lines
    section_keys = ['summary', 'skills', 'experience', 'education']
    section_weights = [priorities.get(k, default_priorities[k]) for k in section_keys]
    total_weight = sum(section_weights)
    section_alloc = {k: max(1, math.floor(total_lines * w / total_weight)) for k, w in zip(section_keys, section_weights)}
    # Within education, allocate among degrees, specializations, certificates
    edu_total = section_alloc['education']
    edu_keys = ['degrees', 'specializations', 'certificates']
    edu_weights = [priorities.get(k, default_priorities[k]) for k in edu_keys]
    edu_weight_sum = sum(edu_weights)
    edu_alloc = {k: max(0, math.floor(edu_total * w / edu_weight_sum)) for k, w in zip(edu_keys, edu_weights)}
    section_alloc.update(edu_alloc)
    return section_alloc

# --- New: UI for line allocation override ---
def line_allocation_ui(section_alloc):
    st.subheader("Line Allocation (One-Page Limit)")
    st.write("Adjust the number of lines for each section. Total must not exceed 60.")
    total_lines = 60
    new_alloc = {}
    used = 0
    for k, v in section_alloc.items():
        if k in ['degrees', 'specializations', 'certificates']:
            label = f"Education: {k.capitalize()}"
        else:
            label = k.capitalize()
        new_alloc[k] = st.number_input(label, min_value=0, max_value=total_lines, value=int(v), step=1, key=f"alloc_{k}_input")
        used += new_alloc[k]
    st.write(f"Total lines used: {used} / {total_lines}")
    if used > total_lines:
        st.error("Total lines exceed 60! Please reduce some sections.")
    return new_alloc, used

# --- Insert before resume generation ---
if 'resume_json_data' in st.session_state and st.session_state.resume_json_data:
    st.header("✏️ Edit Your Resume Sections")
    resume_data = st.session_state.resume_json_data
    # --- New: Analyze job post and allocate lines ---
    if 'section_alloc' not in st.session_state:
        st.session_state.section_alloc = analyze_job_post_and_allocate_lines(
            st.session_state.job_description, resume_data, total_lines=60)
    # --- New: UI for user override ---
    st.session_state.section_alloc, used_lines = line_allocation_ui(st.session_state.section_alloc)
    # Header
    st.subheader("Header")
    header = resume_data.get("header", {})
    header["name"] = st.text_input("Name", value=header.get("name", ""), key="header_name_edit_main")
    header["email"] = st.text_input("Email", value=header.get("email", ""), key="header_email_edit_main")
    header["phone"] = st.text_input("Phone", value=header.get("phone", ""), key="header_phone_edit_main")
    header["linkedin"] = st.text_input("LinkedIn", value=header.get("linkedin", ""), key="header_linkedin_edit_main")
    header["github"] = st.text_input("Github", value=header.get("github", ""), key="header_github_edit_main")
    header["website"] = st.text_input("Website", value=header.get("website", ""), key="header_website_edit_main")
    # Summary
    st.subheader("Professional Summary")
    resume_data["summary"] = st.text_area("Summary", value=resume_data.get("summary", ""), height=80)
    # Add creativity slider and regenerate button
    summary_temp = st.slider("Creativity (Temperature)", min_value=0.0, max_value=1.0, value=0.7, step=0.05, key="summary_temp_input")
    if st.button("Regenerate Professional Summary"):
        # Use current resume data as context
        skills = resume_data.get("skills", [])
        experience = resume_data.get("experience", [])
        education = resume_data.get("education", {})
        # Build a prompt
        prompt = (
            "Given the following resume data, generate a concise, impactful professional summary paragraph. "
            "Highlight the candidate's strengths, relevant experience, and technical skills. Return only the summary paragraph.\n"
            f"Skills: {skills}\nExperience: {experience}\nEducation: {education}"
        )
        try:
            response = openai.chat.completions.create(
                model="gpt-4",
                messages=[{"role": "user", "content": prompt}],
                temperature=summary_temp,
                max_tokens=200,
            )
            new_summary = response.choices[0].message.content.strip()
            resume_data["summary"] = new_summary
            st.success("Professional Summary regenerated!")
            st.rerun()
        except Exception as e:
            st.error(f"Error regenerating summary: {e}")
    # Skills
    st.subheader("Technical Skills")
    skills = resume_data.get("skills", [])
    skills = st.data_editor(skills, key="skills_editor_input", num_rows="dynamic")
    resume_data["skills"] = skills
    # Experience
    st.subheader("Professional Experience")
    experience = resume_data.get("experience", [])
    new_experience = []
    for i, job in enumerate(experience):
        # --- Pre-populate bullet_points from first application if needed (for jobs 2/3) ---
        if (not job.get("bullet_points")) and job.get("applications") and job["applications"] and job["applications"][0].get("title", "") == "":
            job["bullet_points"] = job["applications"][0].get("details", [])
            job["applications"] = job["applications"][1:]
        st.markdown(f"**Job {i+1}**")
        job["position"] = st.text_input(f"Position {i+1}", value=job.get("position", ""), key=f"pos_{i}_input")
        job["company"] = st.text_input(f"Company {i+1}", value=job.get("company", ""), key=f"comp_{i}_input")
        job["location"] = st.text_input(f"Location {i+1}", value=job.get("location", ""), key=f"loc_{i}_input")
        job["date_range"] = st.text_input(f"Date Range {i+1}", value=job.get("date_range", ""), key=f"date_{i}_input")
        # Dynamic bullet points (for jobs 2/3)
        bullet_points = job.get("bullet_points", [])
        if f"add_bullet_{i}_input" not in st.session_state:
            st.session_state[f"add_bullet_{i}_input"] = False
        st.markdown("**Bullet Points**")
        # Generate bullet points button
        if st.button("Generate Bullet Points", key=f"gen_bp_{i}_input"):
            job_prompt = f"Generate 3-5 concise, impactful bullet points describing the key responsibilities and achievements for the following job. Return ONLY a JSON list of strings, no commentary.\nJob Title: {job.get('position','')}\nCompany: {job.get('company','')}\nLocation: {job.get('location','')}\nDate Range: {job.get('date_range','')}"
            try:
                response = openai.chat.completions.create(
                    model="gpt-4",
                    messages=[{"role": "user", "content": job_prompt}],
                    temperature=0.7,
                    max_tokens=300,
                )
                bp_output = response.choices[0].message.content.strip()
                import json as _json
                try:
                    new_bullets = _json.loads(bp_output)
                    if isinstance(new_bullets, list):
                        bullet_points = new_bullets
                        job["bullet_points"] = bullet_points
                        st.session_state[f"bp_regen_{i}_input"] = bullet_points
                        st.success("Bullet points generated!")
                        st.rerun()
                    else:
                        st.error("OpenAI did not return a JSON list. Output was:\n" + bp_output)
                except Exception:
                    st.error("Could not parse bullet points JSON. Output was:\n" + bp_output)
            except Exception as e:
                st.error(f"Error generating bullet points: {e}")
        if f"bp_regen_{i}_input" in st.session_state:
            bullet_points = st.session_state[f"bp_regen_{i}_input"]
            job["bullet_points"] = bullet_points
        for b_idx, bullet in enumerate(bullet_points):
            cols = st.columns([0.85, 0.15])
            bullet_points[b_idx] = cols[0].text_input(f"Bullet Point {b_idx+1} (Job {i+1})", value=bullet, key=f"bp_{i}_input_{b_idx}")
            if cols[1].button("Remove", key=f"remove_bp_{i}_input_{b_idx}"):
                bullet_points.pop(b_idx)
                st.rerun()
        if st.button("Add Bullet Point", key=f"add_bp_{i}_input"):
            bullet_points.append("")
            st.rerun()
        job["bullet_points"] = bullet_points
        # Applications (for job 1 and any other jobs with titled applications)
        applications = job.get("applications", [])
        new_applications = []
        for j, app in enumerate(applications):
            app["title"] = st.text_input(f"Title {i+1}.{j+1}", value=app.get("title", ""), key=f"title_{i}_input_{j}")
            details_str = "\n".join(app.get("details", []))
            details_str = st.text_area(f"Details {i+1}.{j+1} (one bullet per line)", value=details_str, key=f"details_{i}_input_{j}")
            app["details"] = [line.strip() for line in details_str.splitlines() if line.strip()]
            new_applications.append(app)
        job["applications"] = new_applications
        new_experience.append(job)
    resume_data["experience"] = new_experience
    # Education
    st.subheader("Education")
    education = resume_data.get("education", {})
    certificates = education.get("certificates", [])
    certificates = st.data_editor(certificates, key="certificates_editor_input", num_rows="dynamic")
    education["certificates"] = certificates
    specializations = education.get("specializations", [])
    specializations = st.data_editor(specializations, key="specializations_editor_input", num_rows="dynamic")
    education["specializations"] = specializations
    degrees = education.get("degrees", [])
    degrees = st.data_editor(degrees, key="degrees_editor_input", num_rows="dynamic")
    education["degrees"] = degrees
    resume_data["education"] = education
    # Button to generate DOCX from edited data
    if st.button("Generate Resume with Edits"):
        docx_buffer = BytesIO()
        doc_out = build_resume_docx(resume_data)
        doc_out.save(docx_buffer)
        docx_buffer.seek(0)
        st.session_state.resume_docx_buffer = docx_buffer.getvalue()
        # Ensure filename variables are defined here
        name = resume_data.get('header', {}).get('name', 'Resume')
        company = st.session_state.company_name
        today_str = date.today().isoformat()
        resume_filename = f"{name} Resume for {company} ({today_str}).docx"
        st.download_button(
            label="Download Edited Resume (DOCX)",
            data=st.session_state.resume_docx_buffer,
            file_name=resume_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="resume-docx-edited"
        ) 