from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from section_title_template import add_section_title

def add_certificates(doc, certificates):
    # Heading
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run("Professional Certificates:")
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(10)
    run.bold = True
    # Certificates
    for i, cert in enumerate(certificates):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(10) if i == len(certificates) - 1 else Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run1 = p.add_run(cert['name'])
        run1.font.name = 'Times New Roman'
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run1.font.size = Pt(10)
        run1.bold = True
        run2 = p.add_run(" | ")
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(10)
        run3 = p.add_run(cert['date'])
        run3.font.name = 'Times New Roman'
        run3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run3.font.size = Pt(10)
        run3.bold = False
        run3.italic = True

def add_specializations(doc, specializations):
    for spec in specializations:
        # Line 1: Institution, Location (bold)
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after = Pt(0)
        p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run1 = p1.add_run(f"{spec['institution']}, {spec['location']}")
        run1.font.name = 'Times New Roman'
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run1.font.size = Pt(10)
        run1.bold = True
        # Line 2: Specialization name (bold), | Completed (italic)
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(10)
        p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run2a = p2.add_run(spec['specialization'])
        run2a.font.name = 'Times New Roman'
        run2a._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run2a.font.size = Pt(10)
        run2a.bold = True
        run2b = p2.add_run(" | Completed ")
        run2b.font.name = 'Times New Roman'
        run2b.font.size = Pt(10)
        run2c = p2.add_run(spec['date'])
        run2c.font.name = 'Times New Roman'
        run2c._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run2c.font.size = Pt(10)
        run2c.bold = False
        run2c.italic = True

def add_degrees(doc, degrees):
    for deg in degrees:
        # Line 1: University, Location (bold), | Graduated (italic)
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after = Pt(0)
        p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run1 = p1.add_run(f"{deg['university']}, {deg['location']}")
        run1.font.name = 'Times New Roman'
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run1.font.size = Pt(10)
        run1.bold = True
        run2 = p1.add_run(" | Graduated ")
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(10)
        run3 = p1.add_run(deg['date'])
        run3.font.name = 'Times New Roman'
        run3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run3.font.size = Pt(10)
        run3.bold = False
        run3.italic = True
        # Line 2: Degree name in major, Minor(s) in ... (bold)
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(10)
        p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run4 = p2.add_run(deg['degree'])
        run4.font.name = 'Times New Roman'
        run4._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run4.font.size = Pt(10)
        run4.bold = True

# Sample usage
doc = Document()
add_section_title(doc, 'Education')
add_certificates(doc, [
    {'name': 'Vanderbilt University Prompt Engineering', 'date': 'November 2024'},
    {'name': 'IBM Data Analysis and Visualization Foundations', 'date': 'November 2024'}
])
add_specializations(doc, [
    {'institution': 'Purdue University', 'location': 'Online (Partnered with Simplilearn)', 'specialization': 'Applied Generative AI Specialization', 'date': 'October, 2024'}
])
add_degrees(doc, [
    {'university': 'University of Miami', 'location': 'Coral Gables, FL', 'date': 'May 2016', 'degree': 'Bachelor of Arts in Economics, Minors in Motion Pictures, Psychology and French'}
])
doc.save('education_section_sample.docx')
print('Saved as education_section_sample.docx') 