from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from section_title_template import add_section_title
from job_heading_template import add_job_heading
from job_details_template import add_job_applications
from education_section_template import add_certificates, add_specializations, add_degrees
from resume_header_template import add_hyperlink

def add_resume_header(doc, header):
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after = Pt(0)
    p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run1 = p1.add_run(header.get("name", ""))
    run1.font.name = 'Times New Roman'
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run1.font.size = Pt(18)
    run1.bold = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if header.get("email"): add_hyperlink(p2, f"mailto:{header['email']}", header['email'], font_size=10)
    p2.add_run(" ⬥ ").font.size = Pt(10)
    if header.get("phone"): p2.add_run(header["phone"]).font.size = Pt(10)
    p2.add_run(" ⬥ ").font.size = Pt(10)
    if header.get("linkedin"): add_hyperlink(p2, header["linkedin"], "LinkedIn", font_size=10)
    p2.add_run(" ⬥ ").font.size = Pt(10)
    if header.get("github"): add_hyperlink(p2, header["github"], "Github", font_size=10)
    p2.add_run(" ⬥ ").font.size = Pt(10)
    if header.get("website"): add_hyperlink(p2, header["website"], "Personal Website", font_size=10)

def add_professional_summary(doc, summary):
    add_section_title(doc, 'Professional Summary')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.left_indent = Pt(0)
    p.paragraph_format.right_indent = Pt(0)
    run = p.add_run(summary)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(10)

def add_technical_skills(doc, skills):
    add_section_title(doc, 'Technical Skills')
    for i, skill in enumerate(skills):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(10) if i == len(skills) - 1 else Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.left_indent = Pt(0)
        p.paragraph_format.right_indent = Pt(0)
        run_skill = p.add_run(f"{skill['category']}:")
        run_skill.font.name = 'Times New Roman'
        run_skill._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run_skill.font.size = Pt(10)
        run_skill.bold = True
        run_desc = p.add_run(f" {skill['details']}")
        run_desc.font.name = 'Times New Roman'
        run_desc._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run_desc.font.size = Pt(10)
        run_desc.bold = False

def add_professional_experience(doc, experience):
    add_section_title(doc, 'Professional Experience')
    for job in experience:
        add_job_heading(
            doc,
            position=job.get("position"),
            date_range=job.get("date_range"),
            company=job.get("company"),
            location=job.get("location"),
            extra_line=None
        )
        bullet_points = job.get("bullet_points", [])
        applications = job.get("applications", [])
        # Insert bullet points as a pseudo-application at the top
        all_applications = []
        if bullet_points:
            # Each bullet point is a detail, no title
            all_applications.append({"title": "", "details": bullet_points})
        all_applications.extend(applications)
        if all_applications:
            add_job_applications(doc, all_applications)

def add_education(doc, education):
    add_section_title(doc, 'Education')
    if education.get("certificates"):
        add_certificates(doc, education["certificates"])
    if education.get("specializations"):
        add_specializations(doc, education["specializations"])
    if education.get("degrees"):
        add_degrees(doc, education["degrees"])

def build_resume_docx(resume_data):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.gutter = Inches(0)
    add_resume_header(doc, resume_data.get("header", {}))
    add_professional_summary(doc, resume_data.get("summary", ""))
    add_technical_skills(doc, resume_data.get("skills", []))
    add_professional_experience(doc, resume_data.get("experience", []))
    add_education(doc, resume_data.get("education", {}))
    return doc 