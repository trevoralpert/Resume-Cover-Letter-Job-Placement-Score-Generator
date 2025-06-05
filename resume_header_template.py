from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Helper to add a hyperlink to a paragraph
def add_hyperlink(paragraph, url, text, font_name='Times New Roman', font_size=10):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size * 2)))
    rPr.append(sz)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')  # Blue
    rPr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')  # Underline
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph

doc = Document()

# Line 1: Name
p1 = doc.add_paragraph()
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p1.paragraph_format.space_before = Pt(0)
p1.paragraph_format.space_after = Pt(0)
p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
run1 = p1.add_run("Trevor Alpert")
run1.font.name = 'Times New Roman'
run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
run1.font.size = Pt(18)
run1.bold = True

# Line 2: Contact info
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after = Pt(0)
p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

# Email (hyperlinked)
add_hyperlink(p2, 'mailto:trevoralpert1@gmail.com', 'trevoralpert1@gmail.com', font_size=10)

# Diamond bullet
p2.add_run(" ⬥ ").font.size = Pt(10)

# Phone (plain text)
run3 = p2.add_run("(561) 271-9889")
run3.font.name = 'Times New Roman'
run3.font.size = Pt(10)

# Diamond bullet
p2.add_run(" ⬥ ").font.size = Pt(10)

# LinkedIn (hyperlinked)
add_hyperlink(p2, 'https://www.linkedin.com/in/trevoralpert', 'LinkedIn', font_size=10)

# Diamond bullet
p2.add_run(" ⬥ ").font.size = Pt(10)

# Github (hyperlinked)
add_hyperlink(p2, 'https://www.github.com/trevoralpert', 'Github', font_size=10)

# Diamond bullet
p2.add_run(" ⬥ ").font.size = Pt(10)

# Personal Website (hyperlinked)
add_hyperlink(p2, 'https://www.trevoralpert.com', 'Personal Website', font_size=10)

doc.save("resume_header_sample.docx")
print("Saved as resume_header_sample.docx") 