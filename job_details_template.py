from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from job_heading_template import add_job_heading

def add_job_applications(doc, applications):
    for app in applications:
        # Title line (not bulleted, bold, no indent)
        if app['title']:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.left_indent = Pt(0)
            run = p.add_run(app['title'])
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            run.font.size = Pt(10)
            run.bold = True
        # Bulleted details (not bold, indented)
        for i, detail in enumerate(app['details']):
            p_bullet = doc.add_paragraph(style='List Bullet')
            p_bullet.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_bullet.paragraph_format.space_before = Pt(0)
            # Last bullet of last app gets 10pt after, others 0pt
            is_last_bullet = (i == len(app['details']) - 1) and (app == applications[-1])
            p_bullet.paragraph_format.space_after = Pt(10) if is_last_bullet else Pt(0)
            p_bullet.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p_bullet.paragraph_format.left_indent = Inches(0.25)
            p_bullet.paragraph_format.right_indent = Pt(0)
            run_bullet = p_bullet.add_run(detail)
            run_bullet.font.name = 'Times New Roman'
            run_bullet._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            run_bullet.font.size = Pt(10)
            run_bullet.bold = False

# Sample usage
doc = Document()
# Most recent job (detailed applications)
add_job_heading(doc, position="LLM Prompt Engineer (Freelance)", date_range="October 2024 – Present", company=None, location=None, extra_line="Applications Developed:")
most_recent_applications = [
    {
        'title': "Property Listing Spreadsheet Generator | Client: APM Real Estate",
        'details': [
            "Integrates Selenium for dynamic web scraping and Pinecone and Huggingface for vector embeddings allowing user to input a zip code to generate a spreadsheet with the contact information of the property managers of listings in the area."
        ]
    },
    {
        'title': "AI-Powered HR Assistant | Client: Vytrac Health, Inc.",
        'details': [
            "A chatbot that uses RAG to intake Corporate HR handbooks and answer policy questions."
        ]
    },
    {
        'title': "HR Resume Screening Assistant | Client: Vytrac Health, Inc.",
        'details': [
            "Uses Pinecone vector embeddings to analyze multiple resumes and return any number of Top K matches for the criteria entered in the job description field."
        ]
    },
    {
        'title': "Multiple Choice Question Creator App | Client: Culver City School Board",
        'details': [
            "Intakes a corpus of information students will be tested on, generates embeddings, stores vectors in Pinecone, retrieves relevant docs using OpenAI and Huggingface models and generates Multiple Choice Question Quiz."
        ]
    },
    {
        'title': "YouTube Script Writing Tool | Client: Mythical Kitchen",
        'details': [
            "Uses LangChain, Streamlit UI, Pinecone, OpenAI, and LLAMA 2 to generate scripts for YouTube videos based on desired topic, video length, and sliding scale level of creativity."
        ]
    }
]
add_job_applications(doc, most_recent_applications)

# Older job (bullet points only)
add_job_heading(doc, position="Writers' Assistant", date_range="August 2017 - March 2020", company="ABC Studios", location="Los Angeles, CA", extra_line=None)
older_job_details = [
    "Organized and maintained data for writers' room scripts and notes, ensuring accuracy and easy access for team members.",
    "Assisted in the creation and management of data reports to keep track of project timelines and production requirements."
]
# For older jobs, treat each bullet as a single-item application for formatting (no title)
add_job_applications(doc, [{'title': '', 'details': [d]} for d in older_job_details])

doc.save('job_details_sample.docx')
print('Saved as job_details_sample.docx') 