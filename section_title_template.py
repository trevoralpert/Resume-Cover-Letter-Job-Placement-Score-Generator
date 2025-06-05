from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_section_title(doc, title_text):
    # Add the section title paragraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title_text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(13)
    run.bold = True
    run.italic = True
    # Add a bottom border (horizontal line) to the paragraph
    p_format = p.paragraph_format
    p_format.space_after = Pt(0)
    p_format.space_before = Pt(0)
    p_format.line_spacing = 1.0
    p_elm = p._element
    pPr = p_elm.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')  # thickness
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pbdr.append(bottom)
    pPr.append(pbdr)

# Sample usage
doc = Document()
add_section_title(doc, 'Professional Summary')
doc.save('section_title_sample.docx')
print('Saved as section_title_sample.docx') 