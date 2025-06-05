from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

def add_job_heading(doc, position, date_range, company=None, location=None, extra_line=None):
    # Line 1: Company, Location (optional)
    if company and location:
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after = Pt(0)
        p1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run1 = p1.add_run(f"{company}, {location}")
        run1.font.name = 'Times New Roman'
        run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run1.font.size = Pt(10)
        run1.bold = True
    # Line 2: Position | Date Range
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run2a = p2.add_run(position)
    run2a.font.name = 'Times New Roman'
    run2a._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run2a.font.size = Pt(10)
    run2a.bold = True
    run2b = p2.add_run(" | ")
    run2b.font.name = 'Times New Roman'
    run2b.font.size = Pt(10)
    run2c = p2.add_run(date_range)
    run2c.font.name = 'Times New Roman'
    run2c._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run2c.font.size = Pt(10)
    run2c.italic = True
    # Line 3: Extra heading line (optional)
    if extra_line:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after = Pt(5)
        p3.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run3 = p3.add_run(extra_line)
        run3.font.name = 'Times New Roman'
        run3._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run3.font.size = Pt(10)
        run3.bold = True

# Sample usage
doc = Document()
# Freelance job (no company/location, with extra line)
add_job_heading(doc, position="LLM Prompt Engineer (Freelance)", date_range="October 2024 – Present", company=None, location=None, extra_line="Applications Developed:")
# Standard job (with company/location, no extra line)
add_job_heading(doc, position="Writers' Assistant", date_range="August 2017 - March 2020", company="ABC Studios", location="Los Angeles, CA", extra_line=None)
doc.save('job_heading_sample.docx')
print('Saved as job_heading_sample.docx') 